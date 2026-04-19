"""
Generate coursework technical report as a .docx (edit in Microsoft Word).
Run from repo root: python scripts/build_technical_report_docx.py

This version is organised around TECHNOLOGY STACK CHOICES (rationale, trade-offs,
alternatives rejected). Trim to ~5 pages for submission if required.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_path = root / "docs" / "XJCO3011_Technical_Report.docx"

    doc = Document()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("Technical Report")
    r.bold = True
    r.font.size = Pt(16)

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run(
        "XJCO3011 — Web Services and Web Data\n"
        "Coursework 1: Global Air Quality Anomaly & City Comparison API"
    )
    sr.font.size = Pt(12)

    doc.add_paragraph()
    add_para(
        doc,
        "Student name: [YOUR NAME]    Student ID: [YOUR ID]    Date: April 2026",
        italic=True,
    )
    add_para(
        doc,
        "Focus of this report: why each layer of the stack was chosen, what alternatives "
        "were considered, and what consequences follow for data modelling, ingestion, and "
        "analytics. The module brief limits length (~5 pages); shorten subsections if needed.",
        italic=True,
    )
    doc.add_paragraph()

    # --- 1 Short context (not the main focus) ---
    add_heading(doc, "1. Project aim (one paragraph)", 1)
    add_para(
        doc,
        "The API integrates open air-quality time series with open hourly weather data so "
        "users can compare cities, inspect daily trends, and run simple anomaly screening. "
        "The rest of this document explains how that goal drove specific technology choices "
        "rather than listing features.",
    )

    # --- 2 STACK OVERVIEW ---
    add_heading(doc, "2. Technology stack — overview", 1)
    add_para(
        doc,
        "The implementation is intentionally a small, explicit stack so that design "
        "decisions remain visible in the repository:",
    )
    add_bullets(
        doc,
        [
            "Runtime: Python 3",
            "Web/API: FastAPI on Uvicorn (ASGI)",
            "Contracts & validation: Pydantic v2 (+ pydantic-settings for configuration)",
            "ORM & persistence: SQLAlchemy 2.0 declarative models",
            "Database: SQLite for frictionless local development; PostgreSQL recommended when deployed",
            "Outbound HTTP for ETL: httpx (sync client in scripts; suitable for batch jobs)",
            "Configuration: python-dotenv loading a local .env (secrets not committed)",
            "External data: OpenAQ v3 (air) + Open-Meteo Archive (weather grid)",
        ],
    )

    # --- 3 PYTHON ---
    add_heading(doc, "3. Why Python as the implementation language", 1)
    add_para(
        doc,
        "Python is not the only viable language for a data-centric API, but it offers the "
        "strongest balance for this coursework: mature libraries for HTTP and dates, "
        "excellent interoperability with scientific workflows if the project were extended, "
        "and readable syntax that keeps the report aligned with the actual code. "
        "Alternatives such as Go or Rust would offer lower latency and static binaries, but "
        "would slow iteration for exploratory data integration and would add friction for "
        "coursework markers who expect to run `pip install` and `uvicorn` with minimal ceremony.",
    )
    add_para(
        doc,
        "Type hints (used with SQLAlchemy Mapped[] and FastAPI dependencies) document intent "
        "without requiring a separate compilation step. That choice favours maintainability "
        "over micro-optimisation, which matches an academic assessment rather than a "
        "high-throughput production service.",
    )

    # --- 4 FASTAPI ---
    add_heading(doc, "4. Web framework: FastAPI and why not Django or Flask", 1)
    add_para(
        doc,
        "Django would be appropriate if the coursework required server-rendered pages, admin "
        "CRUD for staff users, or a built-in user model. This project is strictly an API with "
        "JSON payloads; carrying Django’s ORM, settings module, and middleware stack would add "
        "conceptual surface area without being exercised. FastAPI instead optimises for "
        "API-first development: route functions return plain models or dicts, dependency "
        "injection wires database sessions per request, and the same type annotations feed "
        "both validation and OpenAPI schema generation.",
    )
    add_para(
        doc,
        "Compared to Flask, FastAPI provides structured validation and automatic documentation "
        "by default. Flask can achieve parity with extensions (Marshmallow, flasgger, etc.), "
        "but that reintroduces integration decisions the brief does not reward. FastAPI’s "
        "native OpenAPI output directly supports the coursework requirement to supply API "
        "documentation and gives examiners an executable contract at /openapi.json.",
    )
    add_para(
        doc,
        "The coursework also mentioned GraphQL as an allowed alternative. GraphQL excels when "
        "clients need flexible field selection over a graph. Here, consumers are primarily "
        "doing fixed aggregations and CRUD slices; REST with explicit query parameters keeps "
        "caching semantics simple and avoids N+1 resolver complexity for a solo developer. "
        "REST + OpenAPI was therefore the pragmatic choice.",
    )
    add_para(
        doc,
        "FastAPI builds on Starlette for networking primitives (routing, middleware, background "
        "tasks). That inheritance matters when reasoning about performance: the framework is "
        "async-capable, though this codebase uses synchronous SQLAlchemy sessions in request "
        "handlers for clarity. A future optimisation path would be async SQLAlchemy drivers "
        "and async session scopes; it was not adopted here because the dominant cost is "
        "database I/O and external API latency, not Python event-loop overhead, and because "
        "mixing async incorrectly is a common source of subtle bugs for coursework timelines.",
    )
    add_para(
        doc,
        "Interactive documentation (Swagger UI at /docs) is not merely cosmetic: it encodes "
        "the same Pydantic schemas as the runtime validator, so the visible form fields and "
        "example payloads stay aligned with what the server accepts. That reduces the need "
        "for a separate manually maintained Postman collection and supports the module’s "
        "emphasis on web service literacy.",
    )

    # --- 5 PYDANTIC ---
    add_heading(doc, "5. Pydantic, settings, and API contracts", 1)
    add_para(
        doc,
        "Request and response bodies are expressed as Pydantic models. That binds validation "
        "rules to the API boundary: invalid types or missing required fields become HTTP 422 "
        "with structured errors instead of leaking database exceptions. It also encodes a "
        "design stance: the HTTP layer is the authority on what a client may send; the ORM "
        "layer maps only validated shapes into rows.",
    )
    add_para(
        doc,
        "pydantic-settings loads DATABASE_URL and OPENAQ_API_KEY from environment variables "
        "(typically via .env locally). That separates configuration from code and mirrors "
        "twelve-factor style deployment, where the same container image receives different "
        "secrets per environment. A simpler alternative would be hard-coded constants, but "
        "that would undermine reproducibility and would leak credentials if committed.",
    )
    add_para(
        doc,
        "Response models declared on route decorators (response_model=…) constrain serialised "
        "JSON: internal ORM objects may carry relationships or lazy loaders that should not "
        "escape to the wire. Pydantic’s model_dump and FastAPI’s serialisation path therefore "
        "act as an explicit anti-corruption layer between persistence and public API shape.",
    )

    # --- 6 SQLALCHEMY + SQL ---
    add_heading(doc, "6. Persistence: SQLAlchemy and a relational SQL database", 1)
    add_para(
        doc,
        "The brief steers students toward SQL databases unless NoSQL is strongly justified. "
        "Air-quality and weather series are naturally tabular: fixed columns, ordered time "
        "index, foreign keys to cities, and uniqueness constraints for deduplication. A "
        "document store could ingest blobs flexibly, but analytical queries (group-by day, "
        "multi-city filters) map cleanly to SQL and remain inspectable in one query plan.",
    )
    add_para(
        doc,
        "SQLAlchemy 2.0 declarative style was chosen over raw SQL strings to reduce coupling "
        "between Python attribute names and database columns while still allowing explicit "
        "select()/where() constructs in analytics where clarity matters. An alternative is "
        "SQLModel (FastAPI author’s layer); it was not required here because the coursework "
        "already benefits from separate explicit schema classes for OpenAPI documentation.",
    )
    add_para(
        doc,
        "SQLite is the default developer database because it is embedded: no separate server "
        "process, single file, fast feedback loop. Its write locking model is a known "
        "weakness under concurrent writers; PostgreSQL is therefore the recommended "
        "production target when the API is exposed publicly. The same SQLAlchemy engine URL "
        "abstraction makes that switch a configuration change rather than a rewrite.",
    )
    add_para(
        doc,
        "Schema evolution in this snapshot relies on Base.metadata.create_all at application "
        "startup. That trades migration discipline (Alembic revisions) for velocity. The "
        "report states this openly: for coursework it is acceptable; for a long-lived "
        "service, Alembic would be added to version schema alongside code.",
    )
    add_para(
        doc,
        "Two parallel fact tables—measurements for pollutants and weather_measurements for "
        "hourly meteorology—were chosen instead of a single wide table with many nullable "
        "columns. That reflects different cardinalities and units (µg/m³ vs °C), different "
        "deduplication keys, and different upstream APIs. A single-table design would force "
        "either sparse columns or lossy encoding of parameter names; the relational model "
        "keeps each ingestion path independent while still joining on city_id and time when "
        "analytics need both signals.",
    )
    add_para(
        doc,
        "Unique constraints at the database level (for example on city plus timestamp for "
        "weather) complement application-level checks. That is a deliberate stack pairing: "
        "SQLite and Postgres both enforce uniqueness, so a retry after a partial import cannot "
        "silently duplicate rows; the API can map integrity errors to HTTP 409 where "
        "appropriate.",
    )

    # --- 7 UVICORN ASGI ---
    add_heading(doc, "7. ASGI server: Uvicorn", 1)
    add_para(
        doc,
        "FastAPI is an ASGI application; Uvicorn is a production-grade ASGI server with "
        "hot reload for development. Alternatives include Hypercorn or Daphne; Uvicorn is "
        "the de facto default in FastAPI tutorials and integrates cleanly with `--reload` "
        "during development. A WSGI-only stack (gunicorn + Flask) would not apply without "
        "changing the framework.",
    )

    # --- 8 HTTPX DOTENV ingestion ---
    add_heading(doc, "8. Ingestion stack: httpx, dotenv, and offline scripts", 1)
    add_para(
        doc,
        "Batch import is implemented as standalone Python scripts rather than as internal "
        "admin endpoints. That choice reflects separation of concerns: the public API "
        "surface remains CRUD and analytics, while potentially long-running or privileged "
        "imports are operator tasks run from a shell. httpx provides a modern HTTP client "
        "with timeouts and clear exceptions; urllib3/requests would also work, but httpx "
        "aligns with FastAPI’s ecosystem and supports HTTP/2 if needed later.",
    )
    add_para(
        doc,
        "OpenAQ requires an API key header; retries wrap transient TLS or network failures "
        "because student networks and public APIs are not always stable. Pagination (limit, "
        "page) prevents single enormous responses. Deduplication keys mirror database "
        "uniqueness intent so reruns are idempotent—an operational requirement more than a "
        "language feature, but one that influenced the choice to keep imports in Python where "
        "logic is easy to patch.",
    )
    add_para(
        doc,
        "An alternative architecture would push ingestion into the database (foreign data "
        "wrappers, scheduled COPY from S3, or dbt pipelines). Those approaches shine at "
        "warehouse scale but add infrastructure the coursework does not require. Python "
        "scripts keep the entire provenance story in one repository: clone, configure env, "
        "run importer, query API—without provisioning a separate orchestration product.",
    )

    # --- 9 DATA VENDORS ---
    add_heading(doc, "9. External data providers and how they constrained the stack", 1)
    add_para(
        doc,
        "OpenAQ v3 returns rich JSON; integrating it favours a language with strong JSON and "
        "datetime handling. Timestamp fields sometimes appear nested under period.datetimeFrom; "
        "defensive parsing is easier in Python than in shell-only pipelines. Open-Meteo returns "
        "hourly grids; the project stores interpolated values at city coordinates. That is "
        "not a stack decision per se, but it explains why latitude and longitude live on the "
        "City model: the weather stack requires coordinates while pollutant ingestion keys "
        "off city identity and sensor metadata.",
    )
    add_para(
        doc,
        "Licences and citations must appear in the final PDF submission; replace bracketed "
        "references with your exact sensors and dataset pages.",
        italic=True,
    )

    # --- 10 APPLICATION SHAPE FROM STACK ---
    add_heading(doc, "10. Application shape following from the stack", 1)
    add_para(
        doc,
        "Routers group endpoints by bounded context: cities, measurements (air quality), "
        "weather-measurements, analytics. That mirrors how FastAPI encourages modular "
        "include_router composition and keeps OpenAPI tags readable. Analytics are read-only "
        "and implemented with SQLAlchemy select() plus Python statistics for z-scores where "
        "SQL would be less expressive; daily aggregation uses SQL group-by on date() to push "
        "work to the database—an intentional hybrid.",
    )
    add_para(
        doc,
        "Dependency-injected database sessions (typically yield-based generators) scope the "
        "ORM session to one HTTP request. That pattern is idiomatic in FastAPI tutorials and "
        "avoids global connection state. It also clarifies transaction boundaries: commit or "
        "rollback happens predictably after the handler returns, which matters when CRUD "
        "routes and analytics share the same engine but different consistency expectations.",
    )
    add_para(
        doc,
        "HTTP status codes follow conventional REST semantics (200/201/204 for success, "
        "400/404/409/422 for client errors). That choice is partly stylistic, but it also "
        "reflects FastAPI’s tight coupling between validation failures and 422 Unprocessable "
        "Entity—another reason Pydantic sits at the centre of the stack rather than as an "
        "optional add-on.",
    )
    add_para(
        doc,
        "Swagger tags were renamed to distinguish “Air quality — measurements” from "
        "“Weather — hourly” so that interactive docs do not conflate pollutant series with "
        "temperature series. Small UX decisions like that reduce support burden and mirror "
        "how production APIs version and label resources for integrators.",
    )
    add_para(
        doc,
        "Full endpoint contracts and examples remain in the separate API documentation PDF "
        "and in the generated OpenAPI schema; this report does not duplicate every path.",
    )

    # --- 11 TRADE OFFS ---
    add_heading(doc, "11. Trade-offs, limitations, and risks of this stack", 1)
    add_bullets(
        doc,
        [
            "No built-in admin UI (FastAPI vs Django): faster API iteration, less operational tooling.",
            "No authentication layer yet: acceptable for local demo; not acceptable for public production.",
            "SQLite write concurrency: fine for single-user coursework; Postgres needed for real traffic.",
            "create_all without Alembic: faster bootstrap; weaker long-term schema governance.",
            "Synchronous httpx in scripts: simple, but large imports block until complete; a queue "
            "(Celery/RQ) would be the next step if imports became scheduled jobs.",
        ],
    )

    # --- 12 GenAI short ---
    add_heading(doc, "12. Generative AI declaration (fill in)", 1)
    add_para(
        doc,
        "Replace with your tools, purposes, verification steps, and attach conversation logs "
        "as required by the module.",
        italic=True,
    )

    # --- 13 References ---
    add_heading(doc, "13. References", 1)
    add_para(
        doc,
        "FastAPI — https://fastapi.tiangolo.com/\n"
        "Uvicorn — https://www.uvicorn.org/\n"
        "Pydantic — https://docs.pydantic.dev/\n"
        "SQLAlchemy — https://docs.sqlalchemy.org/\n"
        "OpenAQ API — https://docs.openaq.org/\n"
        "Open-Meteo Archive — https://archive-api.open-meteo.com/v1/archive\n"
        "httpx — https://www.python-httpx.org/",
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
        print(f"Wrote: {out_path}")
    except PermissionError:
        alt = out_path.with_name(out_path.stem + "_expanded" + out_path.suffix)
        doc.save(alt)
        print(
            f"Could not overwrite (file open?): {out_path}\n"
            f"Wrote instead: {alt}\n"
            "Close Word/IDE and rename, or run again after closing."
        )


if __name__ == "__main__":
    main()
